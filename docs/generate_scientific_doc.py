# -*- coding: utf-8 -*-
"""
Курсовая работа по проекту Arkanoid Web.
Оформление и повествование — по образцу «пример курсача.txt» (УрФУ, РТФ).
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Pt, Cm, Inches

from coursework_config import (
    CITY_YEAR,
    DISCIPLINE,
    MINISTRY_LINES,
    NORM_CONTROL,
    OUTPUT_FILENAME,
    STUDENT,
    SUPERVISOR,
    TOPIC,
)

FONT = "Times New Roman"
SIZE = Pt(14)
SIZE_SMALL = Pt(12)
INDENT = Cm(1.25)
LINE = 1.5


def _pf(p, indent=True, align=None, before=0, after=0):
    pf = p.paragraph_format
    pf.line_spacing = LINE
    pf.first_line_indent = INDENT if indent else Cm(0)
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if align is not None:
        p.alignment = align


def _run(p, text, bold=False, size=None):
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = size or SIZE
    r.bold = bold
    return r


def body(doc, text, indent=True):
    p = doc.add_paragraph()
    _pf(p, indent=indent)
    _run(p, text)
    return p


def center(doc, text, bold=False, before=0, after=0):
    p = doc.add_paragraph()
    _pf(p, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, before=before, after=after)
    _run(p, text, bold=bold)
    return p


def caps_section(doc, title):
    """ВВЕДЕНИЕ, ЗАКЛЮЧЕНИЕ — прописные, без отступа."""
    p = doc.add_paragraph()
    _pf(p, indent=False, before=12, after=6)
    _run(p, title.upper(), bold=True)
    return p


def chapter(doc, num, title):
    """1 ОБЩЕЕ ОПИСАНИЕ"""
    p = doc.add_paragraph()
    _pf(p, indent=False, before=12, after=6)
    _run(p, f"{num} {title.upper()}", bold=True)
    return p


def subsection(doc, num, title):
    """1.1 История..."""
    p = doc.add_paragraph()
    _pf(p, indent=False, before=6, after=3)
    _run(p, f"{num} {title}", bold=True)
    return p


def subsubsection(doc, num, title):
    """1.3.1 Алгоритм..."""
    p = doc.add_paragraph()
    _pf(p, indent=False, before=3, after=3)
    _run(p, f"{num} {title}", bold=True)
    return p


def table_caption(doc, text):
    p = doc.add_paragraph()
    _pf(p, indent=False, before=6, after=3)
    _run(p, text)
    return p


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = FONT
                    r.font.size = SIZE_SMALL
    doc.add_paragraph()
    return t


def toc_line(doc, title, page):
    p = doc.add_paragraph()
    _pf(p, indent=False)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.2), WD_TAB_ALIGNMENT.RIGHT)
    _run(p, title)
    r_tab = p.add_run()
    r_tab.add_tab()
    r_tab.font.name = FONT
    r_tab.font.size = SIZE
    _run(p, str(page))
    return p


def title_page(doc):
    for line in MINISTRY_LINES:
        if line:
            center(doc, line)
        else:
            doc.add_paragraph()
    for _ in range(4):
        doc.add_paragraph()
    center(doc, "Курсовая работа")
    center(doc, f"по дисциплине {DISCIPLINE}")
    center(doc, "")
    center(doc, f"Тема: {TOPIC}", before=6)
    for _ in range(6):
        doc.add_paragraph()
    for line in (SUPERVISOR, NORM_CONTROL, STUDENT):
        p = doc.add_paragraph()
        _pf(p, indent=False)
        parts = line.split("\t")
        if len(parts) == 2:
            _run(p, parts[0] + "\t")
            _run(p, parts[1])
        else:
            _run(p, line)
    for _ in range(5):
        doc.add_paragraph()
    center(doc, CITY_YEAR, before=24)
    doc.add_page_break()


def abbreviations(doc):
    caps_section(doc, "Обозначения и сокращения")
    body(
        doc,
        "В представленной курсовой работе применяют следующие обозначения и сокращения:",
        indent=False,
    )
    terms = [
        ("Arkanoid", "аркадная видеоигра, в которой игрок управляет платформой и отбивает мяч, разрушая блоки."),
        ("Arkanoid Web", "веб-приложение, реализующее игру Arkanoid в браузере с поддержкой сетевых режимов."),
        ("HTTP (Hypertext Transfer Protocol)", "протокол передачи гипертекста; используется для загрузки страницы и API рекордов."),
        ("HTTPS", "защищённая версия HTTP с шифрованием TLS."),
        ("WebSocket", "протокол полнодуплексного обмена данными поверх TCP; применяется для лобби и игровой сессии."),
        ("ws / wss", "схемы URI для WebSocket без шифрования и с TLS соответственно."),
        ("Apache HTTP Server", "веб-сервер с открытым исходным кодом; раздаёт статику и проксирует WebSocket."),
        ("PHP", "язык сценариев для серверной обработки HTTP-запросов; реализует API таблицы рекордов."),
        ("Node.js", "среда выполнения JavaScript на сервере; обслуживает WebSocket и игровую логику."),
        ("Canvas API", "интерфейс HTML5 для программной отрисовки двумерной графики в браузере."),
        ("JSON (JavaScript Object Notation)", "текстовый формат обмена структурированными данными."),
        ("REST API", "архитектурный стиль взаимодействия через HTTP-методы и ресурсы."),
        ("Solo", "одиночный режим игры без второго участника."),
        ("Duo (coop)", "кооперативный режим: два игрока совместно против блоков."),
        ("Versus", "соревновательный режим: два игрока на противоположных сторонах поля, у каждого свой мяч."),
        ("roomId", "краткий буквенно-цифровой код комнаты для подключения второго игрока."),
        ("systemd", "система инициализации и управления службами в Linux."),
        ("Debian", "дистрибутив операционной системы GNU/Linux."),
        ("ПО", "программное обеспечение."),
        ("ОС", "операционная система."),
        ("API (Application Programming Interface)", "программный интерфейс приложения."),
        ("UI (User Interface)", "пользовательский интерфейс."),
    ]
    for term, definition in terms:
        body(doc, f"{term} — {definition}")
    doc.add_page_break()


def introduction(doc):
    caps_section(doc, "Введение")
    body(
        doc,
        "Развитие веб-технологий и повсеместное распространение браузеров как универсальной "
        "среды выполнения приложений привели к тому, что интерактивные игры всё чаще "
        "реализуются без установки отдельного клиентского ПО. Пользователь получает доступ "
        "к игре по ссылке, а разработчик может централизованно обновлять логику на сервере. "
        "Для многопользовательских режимов при этом необходим канал обмена данными в реальном "
        "времени с низкой задержкой.",
    )
    body(
        doc,
        "Классическая игра Arkanoid относится к жанру аркад с элементами физической симуляции "
        "отскока мяча. Перенос механики в веб-среду требует решения задач доставки статических "
        "ресурсов, организации сессий, согласованного состояния у нескольких клиентов и "
        "сохранения результатов. В проекте Arkanoid Web реализованы режимы Solo (один игрок), "
        "Duo (кооператив) и Versus (соревнование один на один с раздельными мячами и стенками "
        "у границ поля).",
    )
    body(
        doc,
        "Цель работы — разработка и описание веб-приложения Arkanoid Web с развёртыванием "
        "на сервере Debian 12, единой точкой сетевого доступа по портам 80/443 и описанием "
        "взаимодействия пользователя с главным меню и лобби.",
    )
    body(doc, "Задачи работы:", indent=False)
    tasks = [
        "изучить предметную область веб-аркад и протокол WebSocket;",
        "спроектировать архитектуру с разделением Apache, PHP и Node.js;",
        "описать структуру каталогов проекта и назначение компонентов;",
        "рассмотреть сценарии главного меню, лобби и HTTP-запросы таблицы рекордов;",
        "реализовать и описать развёртывание с проксированием WebSocket на путь /ws;",
        "обосновать выбор системных пакетов Debian 12.",
    ]
    for i, t in enumerate(tasks, 1):
        p = doc.add_paragraph()
        _pf(p, indent=False)
        _run(p, f"{i}) {t};")

    body(
        doc,
        "Актуальность работы обусловлена следующими факторами. Во-первых, навыки развёртывания "
        "полноценного веб-приложения с компонентом реального времени востребованы при разработке "
        "интерактивных сервисов и учебных стендов. Во-вторых, разделение статики, REST API и "
        "WebSocket отражает распространённую промышленную практику. В-третьих, использование "
        "только открытого ПО (Apache, PHP, Node.js, Debian) снижает стоимость внедрения.",
    )
    body(
        doc,
        "Практическая значимость работы состоит в том, что описанный программный комплекс "
        "может быть развёрнут на учебном или производственном сервере, использован для "
        "демонстрации сетевых режимов игры и дальнейшего расширения функциональности "
        "(аутентификация, рейтинги, HTTPS).",
    )
    doc.add_page_break()


def build_document():
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = SIZE
    st.paragraph_format.line_spacing = LINE
    st.paragraph_format.first_line_indent = INDENT

    title_page(doc)

    caps_section(doc, "СОДЕРЖАНИЕ")
    toc = [
        ("ОБОЗНАЧЕНИЯ И СОКРАЩЕНИЯ", 4),
        ("ВВЕДЕНИЕ", 6),
        ("1 ОБЩЕЕ ОПИСАНИЕ", 8),
        ("1.1 Предметная область и аналоги", 8),
        ("1.2 Технологии, применённые при разработке", 9),
        ("1.3 Алгоритмы, применённые при разработке", 10),
        ("1.3.1 Централизованный игровой цикл на сервере", 10),
        ("1.3.2 Размещение блоков и стенок в режиме Versus", 11),
        ("1.3.3 Управление комнатами и обратный отсчёт", 12),
        ("2 ТЕХНОЛОГИИ, ПРИМЕНЯЕМЫЕ ПРИ РАЗРАБОТКЕ ПРОГРАММ", 14),
        ("2.1 Сравнение подходов к организации многопользовательской игры", 14),
        ("2.2 Сравнение HTTP и WebSocket для лобби", 15),
        ("2.3 Единый порт доступа: прокси Apache", 16),
        ("3 РАЗРАБОТКА ПРОГРАММЫ", 18),
        ("3.1 Постановка задачи", 18),
        ("3.2 Описание решения поставленной задачи", 19),
        ("3.3 Описание работы программы и её настройки", 20),
        ("3.3.1 Настройка веб-части Apache", 20),
        ("3.3.2 Настройка прокси WebSocket", 21),
        ("3.3.3 Настройка игрового сервиса Node.js", 22),
        ("4 РЕАЛИЗАЦИЯ ПРОГРАММЫ", 23),
        ("4.1 Выбранные технологии для реализации программы", 23),
        ("4.2 Разработанные интерфейсы", 24),
        ("4.3 Описание реализованных модулей", 25),
        ("4.3.1 Клиентский модуль game.js", 25),
        ("4.3.2 Серверный модуль roomManager.js", 26),
        ("ЗАКЛЮЧЕНИЕ", 27),
        ("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 28),
        ("ПРИЛОЖЕНИЕ А Команды развёртывания", 30),
    ]
    for title, page in toc:
        toc_line(doc, title, page)
    doc.add_page_break()

    abbreviations(doc)
    introduction(doc)

    # ——— 1 ОБЩЕЕ ОПИСАНИЕ ———
    chapter(doc, "1", "Общее описание")

    subsection(doc, "1.1", "Предметная область и аналоги")
    body(
        doc,
        "Игра Arkanoid была создана в 1986 году компанией Taito и стала одной из наиболее "
        "узнаваемых аркадных игр. Игрок управляет горизонтальной платформой в нижней части "
        "экрана и отбивает мяч, разрушая массив блоков в верхней части. В веб-реализации "
        "сохраняются базовые механики: движение платформы, отражение мяча от стен и блоков, "
        "подсчёт очков и жизней.",
    )
    body(
        doc,
        "Проект Arkanoid Web расширяет классическую модель тремя режимами. В режиме Solo "
        "участвует один игрок; интерфейс не требует кода комнаты — после нажатия «Начать игру» "
        "устанавливается WebSocket-соединение и отправляется сообщение startSolo. В режиме Duo "
        "два игрока совместно управляют платформами в нижней части поля и используют общий "
        "счёт. В режиме Versus платформы расположены сверху и снизу, у каждого игрока свой "
        "цветной мяч; блоки размещены по центру между платформами; у нижней и верхней границы "
        "экрана расположены цветные стенки, отражающие «свой» мяч и приводящие к потере жизни "
        "при пересечении «чужой» стенки.",
    )

    subsection(doc, "1.2", "Технологии, применённые при разработке")
    body(
        doc,
        "Клиентская часть построена на HTML5, CSS3 и JavaScript без обязательного использования "
        "frontend-фреймворков. Единая страница index.html содержит блоки главного меню, настройки, "
        "лобби ожидания и элемент canvas для отрисовки поля. Файл config.js формирует адрес "
        "WebSocket как ws(s)://<хост>/ws — тот же хост и порт, что у сайта.",
    )
    body(
        doc,
        "Серверная часть включает Apache 2 для раздачи статики и исполнения PHP-скрипта "
        "api/scores.php, а также процесс Node.js с библиотекой ws. Игровой процесс слушает "
        "адрес 127.0.0.1:3001; внешний доступ к нему осуществляется только через модуль "
        "proxy_wstunnel Apache. Управление автозапуском Node выполняется службой systemd "
        "arkanoid-game.",
    )

    subsection(doc, "1.3", "Алгоритмы, применённые при разработке")

    subsubsection(doc, "1.3.1", "Централизованный игровой цикл на сервере")
    body(
        doc,
        "Состояние игры (положение мячей, платформ, блоков, счёт, жизни) хранится и "
        "обновляется на сервере в модуле gameEngine.js с фиксированным интервалом тика "
        "(TICK_MS). Клиенты отправляют только команды ввода (input); сервер рассылает "
        "актуальное состояние (gameState). Такой подход исключает рассинхронизацию при "
        "различной задержке сети у игроков.",
    )
    table_caption(doc, "Таблица 1 – Этапы обработки тика на сервере")
    add_table(
        doc,
        ["№", "Этап", "Описание"],
        [
            ("1", "Применение ввода", "Сдвиг платформы, пуск мяча по запросу игрока"),
            ("2", "Обновление мячей", "Перемещение, отражение от стен и платформ"),
            ("3", "Столкновения", "Обработка блоков, Versus-стенок, потеря жизней"),
            ("4", "Сериализация", "Формирование gameState для клиентов"),
        ],
    )

    subsubsection(doc, "1.3.2", "Размещение блоков и стенок в режиме Versus")
    body(
        doc,
        "Функция createVersusBricks() вычисляет вертикальную полосу между зонами платформ "
        "с отступом VERSUS_PLAY_MARGIN и центрирует сетку блоков 6×10 внутри неё. "
        "Функция handleVersusBoundaryWalls() обрабатывает красную стенку у нижней границы "
        "(отражение красного мяча, потеря жизни синего при пересечении) и синюю стенку у "
        "верхней границы (симметрично для синего мяча).",
    )

    subsubsection(doc, "1.3.3", "Управление комнатами и обратный отсчёт")
    body(
        doc,
        "Модуль roomManager.js генерирует идентификатор комнаты roomId, хранит до двух "
        "подключённых игроков, настройки скорости мяча и запускает таймер countdown "
        "(3, 2, 1 секунды) перед началом матча в режимах Duo и Versus. Для Solo используется "
        "сообщение startSolo без отображения кода комнаты в интерфейсе.",
    )

    doc.add_page_break()

    # ——— 2 ТЕХНОЛОГИИ ———
    chapter(doc, "2", "Технологии, применяемые при разработке программы")

    subsection(doc, "2.1", "Сравнение подходов к организации многопользовательской игры")
    body(
        doc,
        "Многопользовательскую игру в браузере можно реализовать несколькими способами: "
        "периодический опрос сервера (polling), длинные HTTP-запросы (long polling) или "
        "постоянное соединение WebSocket. Для Arkanoid Web выбран WebSocket как обеспечивающий "
        "минимальную задержку и двунаправленную передачу без накладных расходов повторного "
        "установления соединения.",
    )
    table_caption(doc, "Таблица 2 – Сравнение моделей обмена данными")
    add_table(
        doc,
        ["Критерий", "HTTP polling", "WebSocket"],
        [
            ("Задержка", "Равна интервалу опроса", "Минимальная"),
            ("Нагрузка на сервер", "Постоянные запросы", "Одно соединение на клиента"),
            ("Двунаправленность", "Ограничена", "Полная"),
            ("Применение в проекте", "Только рекорды (GET/POST)", "Лобби и игра"),
        ],
    )
    body(
        doc,
        "На основании приведённого сравнения для фазы лобби и игровой сессии выбран WebSocket; "
        "для редких операций с таблицей рекордов достаточно HTTP.",
    )

    subsection(doc, "2.2", "Сравнение HTTP и WebSocket для лобби")
    body(
        doc,
        "Главное меню и лобби используют WebSocket-сообщения с полем type в формате JSON: "
        "startSolo, createRoom, joinRoom, updateSettings. Ответы сервера: soloStarted, "
        "roomCreated, roomJoined, roomUpdate, countdown, error. Таблица рекордов загружается "
        "отдельно через GET /api/scores.php при открытии страницы и не связана с комнатами.",
    )

    subsection(doc, "2.3", "Единый порт доступа: прокси Apache")
    body(
        doc,
        "Для упрощения настройки firewall и соответствия корпоративным политикам внешний "
        "доступ организован только по портам 80 и 443. Модули proxy, proxy_http и proxy_wstunnel "
        "перенаправляют путь /ws на ws://127.0.0.1:3001/. Команды a2enmod и a2enconf "
        "выполняются из любой директории; дополнительные пакеты сверх apache2 не требуются.",
    )

    doc.add_page_break()

    # ——— 3 РАЗРАБОТКА ———
    chapter(doc, "3", "Разработка программы")

    subsection(doc, "3.1", "Постановка задачи")
    body(
        doc,
        "Требуется разработать веб-приложение Arkanoid с режимами Solo, Duo и Versus, "
        "таблицей рекордов и развёртыванием на Debian 12.",
    )
    body(doc, "Входные данные:", indent=False)
    body(
        doc,
        "действия пользователя в интерфейсе (выбор режима, имя, код комнаты, скорость мяча); "
        "команды управления в фазе игры (движение платформы, пуск мяча); HTTP-запросы к API рекордов.",
    )
    body(doc, "Выходные данные:", indent=False)
    body(
        doc,
        "отображение игрового поля и интерфейса в браузере; JSON-ответы api/scores.php; "
        "WebSocket-ответы лобби и игрового состояния.",
    )
    body(doc, "Функциональные требования:", indent=False)
    for req in [
        "одиночный старт без кода комнаты;",
        "создание и вход в комнату по коду для Duo и Versus;",
        "быстрый вход с главного экрана;",
        "настройка скорости мяча хостом до старта;",
        "сохранение рекордов после партии;",
        "прокси WebSocket на том же порту, что и сайт.",
    ]:
        p = doc.add_paragraph()
        _pf(p, indent=False)
        _run(p, f"— {req}")

    body(doc, "Нефункциональные требования:", indent=False)
    for req in [
        "развёртывание на Debian 12 с открытым ПО;",
        "автозапуск игрового сервера через systemd;",
        "Node.js слушает только 127.0.0.1:3001;",
        "права записи scores.json для пользователя www-data.",
    ]:
        p = doc.add_paragraph()
        _pf(p, indent=False)
        _run(p, f"— {req}")

    subsection(doc, "3.2", "Описание решения поставленной задачи")
    body(
        doc,
        "Для выполнения требований применена трёхкомпонентная архитектура. Apache отдаёт "
        "статику и PHP, проксирует /ws на Node. PHP обслуживает только рекорды. Node через "
        "roomManager и gameEngine обеспечивает сессии и авторитетную логику игры. "
        "Разделение упрощает сопровождение: обновление клиентских файлов не требует "
        "перезапуска Node, а перезапуск arkanoid-game не затрагивает отдачу HTML.",
    )

    subsection(doc, "3.3", "Описание работы программы и её настройки")

    subsubsection(doc, "3.3.1", "Настройка веб-части Apache")
    body(
        doc,
        "Каталог www копируется в /var/www/html. Назначаются права www-data на data/ и "
        "scores.json. Проверка: curl http://localhost/api/scores.php возвращает JSON-массив.",
    )

    subsubsection(doc, "3.3.2", "Настройка прокси WebSocket")
    body(
        doc,
        "Файл deploy/apache-arkanoid-ws.conf копируется в /etc/apache2/conf-available/arkanoid-ws.conf. "
        "Выполняются команды: sudo a2enmod proxy proxy_http proxy_wstunnel; sudo a2enconf arkanoid-ws; "
        "sudo apache2ctl configtest; sudo systemctl reload apache2. Порт 3001 в firewall наружу не открывается.",
    )
    table_caption(doc, "Таблица 3 – Команды включения прокси Apache")
    add_table(
        doc,
        ["Команда", "Назначение"],
        [
            ("a2enmod proxy proxy_http proxy_wstunnel", "Включение модулей прокси"),
            ("a2enconf arkanoid-ws", "Подключение конфигурации /ws"),
            ("systemctl reload apache2", "Применение настроек без остановки"),
        ],
    )

    subsubsection(doc, "3.3.3", "Настройка игрового сервиса Node.js")
    body(
        doc,
        "Каталог server размещается в /opt/arkanoid-server, выполняется npm install --production. "
        "Unit-файл arkanoid-game.service задаёт HOST=127.0.0.1 и PORT=3001. Проверка: "
        "systemctl status arkanoid-game; в браузере WebSocket подключается к ws://<сервер>/ws.",
    )
    table_caption(doc, "Таблица 4 – Системные пакеты Debian 12")
    add_table(
        doc,
        ["Пакет", "Назначение"],
        [
            ("apache2", "HTTP, статика, прокси /ws"),
            ("libapache2-mod-php", "api/scores.php"),
            ("php-json", "JSON в PHP"),
            ("nodejs", "WebSocket-сервер"),
            ("npm", "зависимость ws"),
        ],
    )

    doc.add_page_break()

    # ——— 4 РЕАЛИЗАЦИЯ ———
    chapter(doc, "4", "Реализация программы")

    subsection(doc, "4.1", "Выбранные технологии для реализации программы")
    body(
        doc,
        "JavaScript (ES6+) — язык клиента и сервера Node.js. HTML5 Canvas — отрисовка поля. "
        "PHP 8.x — API рекордов. Библиотека ws — WebSocket на сервере. systemd — автозапуск "
        "службы arkanoid-game от имени www-data.",
    )

    subsection(doc, "4.2", "Разработанные интерфейсы")
    body(
        doc,
        "Пользовательский интерфейс реализован на одной странице с переключением блоков "
        "modeSelect, lobby, waitingRoom и canvas. Боковая панель содержит подсказки по "
        "управлению, статус сессии и таблицу рекордов.",
    )
    table_caption(doc, "Таблица 5 – Экраны интерфейса и сетевые запросы (меню и лобби)")
    add_table(
        doc,
        ["Экран / действие", "Запрос", "Назначение"],
        [
            ("Загрузка сайта", "GET /api/scores.php", "Рекорды"),
            ("Solo — старт", "WS: startSolo", "Одиночная игра"),
            ("Создать комнату", "WS: createRoom", "Лобби Duo/Versus"),
            ("Войти по коду", "WS: joinRoom", "Второй игрок"),
            ("Смена скорости", "WS: updateSettings", "Хост, до старта"),
        ],
    )
    body(
        doc,
        "Сообщения input, gameState, continue, restart относятся к игровому циклу и в таблицу "
        "не включены в соответствии с границами описания интерфейса лобби.",
    )

    subsection(doc, "4.3", "Описание реализованных модулей")

    subsubsection(doc, "4.3.1", "Клиентский модуль game.js")
    body(
        doc,
        "Модуль управляет переключением экранов (showModeSelect, showLobbyConfig, showWaitingRoom, "
        "showGame), открытием WebSocket через getWebSocketUrl(), обработкой сообщений лобби "
        "в handleServerMessage и загрузкой рекордов loadScores(). Функция returnToMainMenu() "
        "закрывает сокет и сбрасывает sessionActive.",
    )

    subsubsection(doc, "4.3.2", "Серверный модуль roomManager.js")
    body(
        doc,
        "Класс RoomManager реализует createRoom, addPlayer, getRoomSettings, maybeStartCountdown "
        "и startSolo. Хранит комнаты в структуре Map. При подключении двух игроков запускает "
        "обратный отсчёт и передаёт управление gameEngine для старта матча.",
    )
    body(
        doc,
        "Модуль server.js разбирает входящий JSON по полю type и делегирует вызовы RoomManager. "
        "Модуль gameEngine.js содержит createVersusBricks, handleVersusBoundaryWalls, serializeState "
        "и тиковый цикл обновления состояния.",
    )

    doc.add_page_break()

    # ——— ЗАКЛЮЧЕНИЕ ———
    caps_section(doc, "Заключение")
    body(
        doc,
        "В ходе выполнения данной курсовой работы было разработано и описано веб-приложение "
        "Arkanoid Web — многопользовательская реализация классической аркады с режимами Solo, "
        "Duo и Versus.",
    )
    body(
        doc,
        "Клиентская часть обеспечивает выбор режима, настройку скорости мяча, создание и вход "
        "в комнату по коду, быстрый вход с главного экрана и отображение таблицы рекордов "
        "через HTTP API. Серверная часть разделена на Apache с PHP для статики и рекордов "
        "и Node.js для WebSocket-сессий с авторитетной игровой логикой.",
    )
    body(
        doc,
        "Организован доступ по единому порту 80/443 за счёт проксирования WebSocket на путь /ws. "
        "Игровой процесс доступен только на localhost:3001, что соответствует требованиям "
        "безопасности и упрощает настройку межсетевого экрана.",
    )
    body(
        doc,
        "Поставленные в работе задачи решены в полном объёме. Выполнено сравнение подходов "
        "к обмену данными, описаны алгоритмы режима Versus, структура проекта, сценарии лобби "
        "и порядок развёртывания на Debian 12 с командами a2enmod, a2enconf и reload Apache.",
    )

    caps_section(doc, "Список использованных источников")
    refs = [
        "1. Fielding R. Architectural Styles and the Design of Network-based Software Architectures : дис. … докт. наук. — 2000.",
        "2. Fette I., Melnikov A. The WebSocket Protocol. RFC 6455. — IETF, 2011. — Текст : электронный.",
        "3. Arkanoid / Wikipedia. — Текст : электронный // Wikipedia : [сайт]. — URL: https://ru.wikipedia.org/wiki/Arkanoid (дата обращения: 23.05.2026).",
        "4. Debian 12 Bookworm Documentation / Debian Project. — Текст : электронный // Debian : [сайт]. — URL: https://www.debian.org/releases/bookworm/ (дата обращения: 23.05.2026).",
        "5. Apache HTTP Server Documentation / Apache Software Foundation. — Текст : электронный // Apache : [сайт]. — URL: https://httpd.apache.org/docs/2.4/ (дата обращения: 23.05.2026).",
        "6. Node.js Documentation / OpenJS Foundation. — Текст : электронный // Node.js : [сайт]. — URL: https://nodejs.org/docs/ (дата обращения: 23.05.2026).",
        "7. PHP Manual / The PHP Group. — Текст : электронный // PHP : [сайт]. — URL: https://www.php.net/manual/ (дата обращения: 23.05.2026).",
        "8. Canvas API / MDN Web Docs. — Текст : электронный // MDN : [сайт]. — URL: https://developer.mozilla.org/docs/Web/API/Canvas_API (дата обращения: 23.05.2026).",
        "9. WebSocket API / MDN Web Docs. — Текст : электронный // MDN : [сайт]. — URL: https://developer.mozilla.org/docs/Web/API/WebSocket (дата обращения: 23.05.2026).",
        "10. mod_proxy_wstunnel / Apache HTTP Server. — Текст : электронный // Apache : [сайт]. — URL: https://httpd.apache.org/docs/2.4/mod/mod_proxy_wstunnel.html (дата обращения: 23.05.2026).",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        _pf(p, indent=False)
        _run(p, ref)

    doc.add_page_break()

    caps_section(doc, "Приложение А")
    body(doc, "Команды развёртывания программного комплекса", indent=False)
    table_caption(doc, "Таблица А.1 – Основные команды установки")
    add_table(
        doc,
        ["№", "Команда", "Назначение"],
        [
            ("1", "apt install apache2 libapache2-mod-php php-json nodejs npm", "Пакеты"),
            ("2", "cp -r www/* /var/www/html/", "Веб-часть"),
            ("3", "cp apache-arkanoid-ws.conf …/conf-available/", "Прокси /ws"),
            ("4", "a2enmod proxy proxy_http proxy_wstunnel", "Модули"),
            ("5", "a2enconf arkanoid-ws && systemctl reload apache2", "Прокси"),
            ("6", "cp server/* /opt/arkanoid-server/ && npm install", "Node"),
            ("7", "systemctl enable --now arkanoid-game", "Служба"),
        ],
    )

    return doc


if __name__ == "__main__":
    document = build_document()
    document.save(OUTPUT_FILENAME)
    print("Saved:", OUTPUT_FILENAME)
