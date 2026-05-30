# -*- coding: utf-8 -*-
import os
import re
import sys
import zipfile
from pathlib import Path

paths = [
    Path(r"c:\Users\wakai\Downloads\Telegram Desktop\Курсач-фулл.docx"),
    Path.home() / "Downloads" / "Telegram Desktop" / "Курсач-фулл.docx",
]

out = Path(__file__).parent / "example_docx_analysis.txt"

def main():
    path = None
    for p in paths:
        if p.exists():
            path = p
            break
    lines = []
    if not path:
        lines.append("FILE NOT FOUND")
        out.write_text("\n".join(lines), encoding="utf-8")
        return

    lines.append(f"PATH: {path}")
    lines.append(f"SIZE: {path.stat().st_size}")

    try:
        from docx import Document
        doc = Document(str(path))
        lines.append(f"PARAGRAPHS: {len(doc.paragraphs)}")
        for i, p in enumerate(doc.paragraphs[:80]):
            style = p.style.name if p.style else "?"
            align = p.alignment
            text = (p.text or "").strip()[:200]
            if not text and not p.runs:
                continue
            fonts = []
            for r in p.runs[:3]:
                fonts.append(
                    f"bold={r.bold} size={r.font.size} name={r.font.name} "
                    f"italic={r.italic}"
                )
            lines.append(
                f"[{i}] style={style} align={align} | {text} | runs: {fonts}"
            )
        # sections sample from middle
        lines.append("--- MIDDLE SAMPLE ---")
        mid = len(doc.paragraphs) // 2
        for i, p in enumerate(doc.paragraphs[mid : mid + 40]):
            text = (p.text or "").strip()
            if text:
                lines.append(f"[{mid+i}] {p.style.name}: {text[:150]}")
    except Exception as e:
        lines.append(f"python-docx error: {e}")
        z = zipfile.ZipFile(path)
        xml = z.read("word/document.xml").decode("utf-8", errors="replace")
        text = re.sub(r"</w:p>", "\n", xml)
        text = re.sub(r"<[^>]+>", "", text)
        lines.append(text[:12000])

    out.write_text("\n".join(lines), encoding="utf-8")
    print("Wrote", out)


if __name__ == "__main__":
    main()
